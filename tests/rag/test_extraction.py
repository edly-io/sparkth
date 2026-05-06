"""
tests/rag/test_extraction.py

Unit tests for app/rag/extraction.py
Covers: PDF, DOCX, HTML, TXT extraction + public API routing + error handling.
"""

import io
from typing import Any
from unittest.mock import MagicMock, patch

import pytest
from bs4 import Tag

from app.rag.extraction import (
    ExtractionResult,
    _bs_table_to_md,
    _docx_table_to_md,
    _extract_docx,
    _extract_html,
    _extract_txt,
    extract_to_markdown,
)
from app.rag.types import DocType

SIMPLE_HTML = b"""
<html><body>
  <h1>Course Introduction</h1>
  <h2>What you will learn</h2>
  <p>This course covers machine learning fundamentals.</p>
  <ul>
    <li>Supervised learning</li>
    <li>Unsupervised learning</li>
  </ul>
  <table>
    <tr><th>Topic</th><th>Week</th></tr>
    <tr><td>Linear Regression</td><td>1</td></tr>
    <tr><td>Neural Networks</td><td>4</td></tr>
  </table>
</body></html>
"""

MINIMAL_HTML = b"<html><body><p>Hello world</p></body></html>"

TXT_CONTENT = b"# Lecture Notes\n\nSome content here.\n"


class TestExtractionResult:
    def test_fields_are_stored(self) -> None:
        r = ExtractionResult("# Hello", DocType.TXT, "file.txt", page_count=1, warnings=["w"])
        assert r.markdown == "# Hello"
        assert r.doc_type == DocType.TXT
        assert r.source_name == "file.txt"
        assert r.page_count == 1
        assert r.warnings == ["w"]

    def test_warnings_default_to_empty_list(self) -> None:
        r = ExtractionResult("md", DocType.TXT, "f.txt")
        assert r.warnings == []

    def test_page_count_defaults_to_none(self) -> None:
        r = ExtractionResult("md", DocType.TXT, "f.txt")
        assert r.page_count is None


class TestExtractPDF:
    def test_calls_pymupdf4llm_and_returns_result(self) -> None:
        fake_md = "# Chapter 1\n\nContent here.\n-----\n# Chapter 2\n"
        with (
            patch("app.rag.extraction.fitz.open"),
            patch("app.rag.extraction.pymupdf4llm.to_markdown", return_value=fake_md) as mock_to_md,
        ):
            result = extract_to_markdown(b"%PDF-fake", "notes.pdf")

        mock_to_md.assert_called_once()
        assert result.doc_type == DocType.PDF
        assert result.source_name == "notes.pdf"
        assert "\n-----\n" not in result.markdown

    def test_fitz_open_called_with_stream_and_filetype(self) -> None:
        raw = b"%PDF-fake"
        with (
            patch("app.rag.extraction.fitz.open") as mock_fitz,
            patch("app.rag.extraction.pymupdf4llm.to_markdown", return_value="md"),
        ):
            extract_to_markdown(raw, "notes.pdf")

        mock_fitz.assert_called_once_with(stream=raw, filetype="pdf")

    def test_separators_stripped_from_markdown(self) -> None:
        fake_md = "page1\n-----\npage2\n-----\npage3"
        with (
            patch("app.rag.extraction.fitz.open"),
            patch("app.rag.extraction.pymupdf4llm.to_markdown", return_value=fake_md),
        ):
            result = extract_to_markdown(b"%PDF-fake", "book.pdf")

        assert "\n-----\n" not in result.markdown
        assert "page1" in result.markdown
        assert "page2" in result.markdown
        assert "page3" in result.markdown

    def test_page_count_inferred_from_separators(self) -> None:
        fake_md = "page1\n-----\npage2\n-----\npage3"
        with (
            patch("app.rag.extraction.fitz.open"),
            patch("app.rag.extraction.pymupdf4llm.to_markdown", return_value=fake_md),
        ):
            result = extract_to_markdown(b"%PDF-fake", "book.pdf")

        assert result.page_count == 3

    def test_single_page_no_separator(self) -> None:
        fake_md = "just one page, no separator"
        with (
            patch("app.rag.extraction.fitz.open"),
            patch("app.rag.extraction.pymupdf4llm.to_markdown", return_value=fake_md),
        ):
            result = extract_to_markdown(b"%PDF-fake", "one.pdf")

        assert result.page_count == 1


class TestExtractHTML:
    def test_h1_rendered_as_h1(self) -> None:
        result = _extract_html(SIMPLE_HTML, "doc.html")
        assert "# Course Introduction" in result.markdown

    def test_h2_rendered_as_h2(self) -> None:
        result = _extract_html(SIMPLE_HTML, "doc.html")
        assert "## What you will learn" in result.markdown

    def test_paragraph_included(self) -> None:
        result = _extract_html(SIMPLE_HTML, "doc.html")
        assert "This course covers machine learning fundamentals." in result.markdown

    def test_list_items_rendered_with_dash(self) -> None:
        result = _extract_html(SIMPLE_HTML, "doc.html")
        assert "- Supervised learning" in result.markdown
        assert "- Unsupervised learning" in result.markdown

    def test_ordered_list_items_rendered_with_numbers(self) -> None:
        html = b"<html><body><ol><li>First</li><li>Second</li></ol></body></html>"
        result = _extract_html(html, "doc.html")
        assert "1. First" in result.markdown
        assert "2. Second" in result.markdown

    def test_table_rendered_as_gfm(self) -> None:
        result = _extract_html(SIMPLE_HTML, "doc.html")
        assert "| Topic |" in result.markdown
        assert "| ---" in result.markdown
        assert "| Linear Regression |" in result.markdown

    def test_table_warning_emitted(self) -> None:
        result = _extract_html(SIMPLE_HTML, "doc.html")
        assert any("Table detected" in w for w in result.warnings)

    def test_no_warnings_for_table_free_html(self) -> None:
        result = _extract_html(MINIMAL_HTML, "simple.html")
        assert result.warnings == []

    def test_doc_type_is_html(self) -> None:
        result = _extract_html(MINIMAL_HTML, "simple.html")
        assert result.doc_type == DocType.HTML

    def test_empty_body_returns_empty_markdown(self) -> None:
        result = _extract_html(b"<html><body></body></html>", "empty.html")
        assert result.markdown.strip() == ""

    def test_content_inside_div_not_dropped(self) -> None:
        html = b"""
        <html><body>
          <div>
            <h2>Nested Heading</h2>
            <p>Nested paragraph.</p>
          </div>
        </body></html>
        """
        result = _extract_html(html, "nested.html")
        assert "## Nested Heading" in result.markdown
        assert "Nested paragraph." in result.markdown

    def test_deeply_nested_containers_traversed(self) -> None:
        html = b"""
        <html><body>
          <div><section><article>
            <h3>Deep Heading</h3>
            <p>Deep content.</p>
          </article></section></div>
        </body></html>
        """
        result = _extract_html(html, "deep.html")
        assert "### Deep Heading" in result.markdown
        assert "Deep content." in result.markdown

    def test_whitespace_only_body_returns_empty(self) -> None:
        html = b"<html><body>  \n  <p>Clean</p>\n  </body></html>"
        result = _extract_html(html, "ws.html")
        assert result.markdown.count("Clean") == 1

    def test_nested_ul_indented(self) -> None:
        html = b"""
        <html><body>
          <ul>
            <li>Parent
              <ul><li>Child</li></ul>
            </li>
          </ul>
        </body></html>
        """
        result = _extract_html(html, "nested.html")
        assert "- Parent" in result.markdown
        assert "  - Child" in result.markdown

    def test_nested_ol_numbered(self) -> None:
        html = b"""
        <html><body>
          <ol>
            <li>First
              <ol><li>Sub-first</li></ol>
            </li>
          </ol>
        </body></html>
        """
        result = _extract_html(html, "nested_ol.html")
        assert "1. First" in result.markdown
        assert "  1. Sub-first" in result.markdown

    def test_nested_list_parent_text_not_duplicated(self) -> None:
        html = b"""
        <html><body>
          <ul>
            <li>Parent
              <ul><li>Child</li></ul>
            </li>
          </ul>
        </body></html>
        """
        result = _extract_html(html, "dup.html")
        assert result.markdown.count("Parent") == 1
        assert result.markdown.count("Child") == 1

    def test_span_direct_child_of_div_not_dropped(self) -> None:
        html = b"<html><body><div><span>Important span text</span></div></body></html>"
        result = _extract_html(html, "doc.html")
        assert "Important span text" in result.markdown

    def test_anchor_direct_child_of_div_not_dropped(self) -> None:
        html = b'<html><body><div><a href="#">Link text</a></div></body></html>'
        result = _extract_html(html, "doc.html")
        assert "Link text" in result.markdown

    def test_multiple_sibling_spans_all_captured(self) -> None:
        html = b"""
        <html><body>
          <div><span>Alpha</span><span>Beta</span></div>
        </body></html>
        """
        result = _extract_html(html, "doc.html")
        assert "Alpha" in result.markdown
        assert "Beta" in result.markdown

    def test_span_inside_paragraph_captured_via_get_text(self) -> None:
        html = b"<html><body><p>Hello <span>world</span></p></body></html>"
        result = _extract_html(html, "doc.html")
        assert "Hello world" in result.markdown

    def test_void_elements_do_not_produce_blank_lines(self) -> None:
        html = b"""
        <html><body>
          <div><br><hr><img src="x.png"><p>After voids</p></div>
        </body></html>
        """
        result = _extract_html(html, "doc.html")
        assert "After voids" in result.markdown
        for token in ("br", "hr", "img"):
            assert token not in result.markdown

    def test_google_docs_span_wrapped_div_pattern(self) -> None:
        html = b"""
        <html><body>
          <div class="c1"><span class="c2">Lecture title</span></div>
          <div class="c1"><span class="c2">Subtitle line</span></div>
        </body></html>
        """
        result = _extract_html(html, "gdoc.html")
        assert "Lecture title" in result.markdown
        assert "Subtitle line" in result.markdown

    def test_blockquote_content_included(self) -> None:
        html = b"<html><body><blockquote>Quoted text</blockquote></body></html>"
        assert "Quoted text" in _extract_html(html, "doc.html").markdown

    def test_pre_content_included(self) -> None:
        html = b"<html><body><pre>code block</pre></body></html>"
        assert "code block" in _extract_html(html, "doc.html").markdown

    def test_figcaption_content_included(self) -> None:
        html = b"""
        <html><body>
          <figure><img src="chart.png"><figcaption>Figure 1: Results</figcaption></figure>
        </body></html>
        """
        assert "Figure 1: Results" in _extract_html(html, "doc.html").markdown

    def test_text_block_inline_children_captured(self) -> None:
        html = b"""
        <html><body>
          <blockquote>Outer <span>inner span</span> text</blockquote>
        </body></html>
        """
        result = _extract_html(html, "doc.html")
        assert "Outer" in result.markdown
        assert "inner span" in result.markdown


class TestExtractTXT:
    def test_content_passed_through_unchanged(self) -> None:
        result = _extract_txt(TXT_CONTENT, "notes.txt")
        assert result.markdown == TXT_CONTENT.decode("utf-8")

    def test_doc_type_is_txt(self) -> None:
        result = _extract_txt(TXT_CONTENT, "notes.txt")
        assert result.doc_type == DocType.TXT

    def test_invalid_utf8_replaced_not_raised(self) -> None:
        bad_bytes = b"Hello \xff\xfe world"
        result = _extract_txt(bad_bytes, "bad.txt")
        assert "Hello" in result.markdown
        assert "world" in result.markdown

    def test_md_extension_uses_txt_extractor(self) -> None:
        result = extract_to_markdown(TXT_CONTENT, "lecture.md")
        assert result.doc_type == DocType.TXT
        assert result.markdown == TXT_CONTENT.decode("utf-8")


class TestExtractDOCX:
    """
    python-docx cannot create a valid .docx from scratch in memory easily,
    so we build a minimal one using python-docx's API and round-trip it.
    """

    @pytest.fixture()
    def simple_docx_bytes(self) -> bytes:
        from docx import Document

        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_heading("Background", level=2)
        doc.add_paragraph("This is a paragraph.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @pytest.fixture()
    def table_docx_bytes(self) -> bytes:
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header A"
        table.cell(0, 1).text = "Header B"
        table.cell(1, 0).text = "Value 1"
        table.cell(1, 1).text = "Value 2"
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @pytest.fixture()
    def ordered_list_docx_bytes(self) -> bytes:
        from docx import Document

        doc = Document()
        for item in ("First", "Second", "Third"):
            doc.add_paragraph(item, style="List Number")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_h1_rendered(self, simple_docx_bytes: bytes) -> None:
        assert "# Introduction" in _extract_docx(simple_docx_bytes, "doc.docx").markdown

    def test_h2_rendered(self, simple_docx_bytes: bytes) -> None:
        assert "## Background" in _extract_docx(simple_docx_bytes, "doc.docx").markdown

    def test_paragraph_included(self, simple_docx_bytes: bytes) -> None:
        assert "This is a paragraph." in _extract_docx(simple_docx_bytes, "doc.docx").markdown

    def test_doc_type_is_docx(self, simple_docx_bytes: bytes) -> None:
        assert _extract_docx(simple_docx_bytes, "doc.docx").doc_type == DocType.DOCX

    def test_table_rendered_as_gfm(self, table_docx_bytes: bytes) -> None:
        result = _extract_docx(table_docx_bytes, "table.docx")
        assert "| Header A |" in result.markdown
        assert "| ---" in result.markdown
        assert "| Value 1 |" in result.markdown

    def test_table_warning_emitted(self, table_docx_bytes: bytes) -> None:
        result = _extract_docx(table_docx_bytes, "table.docx")
        assert any("Table detected" in w for w in result.warnings)

    def test_ordered_list_rendered_with_numbers(self, ordered_list_docx_bytes: bytes) -> None:
        result = _extract_docx(ordered_list_docx_bytes, "doc.docx")
        assert "1. First" in result.markdown
        assert "2. Second" in result.markdown
        assert "3. Third" in result.markdown

    def test_ordered_list_numbering_resets_between_documents(self, ordered_list_docx_bytes: bytes) -> None:
        result1 = _extract_docx(ordered_list_docx_bytes, "doc1.docx")
        result2 = _extract_docx(ordered_list_docx_bytes, "doc2.docx")
        assert "1. First" in result1.markdown
        assert "1. First" in result2.markdown  # must restart, not continue from 4

    def test_list_prefix_attribute_error_falls_back_to_bullet(self, ordered_list_docx_bytes: bytes) -> None:
        with patch("app.rag.extraction._docx_list_prefix", return_value="-"):
            result = _extract_docx(ordered_list_docx_bytes, "doc.docx")

        assert "First" in result.markdown
        assert "- First" in result.markdown


class TestDocxListPrefix:
    """Unit tests for _docx_list_prefix, called directly to exercise its internals."""

    def test_attribute_error_returns_dash(self) -> None:
        from app.rag.extraction import _docx_list_prefix

        bad_num_pr = MagicMock()
        bad_num_pr.find.side_effect = AttributeError("injected")
        assert _docx_list_prefix(MagicMock(), bad_num_pr, {}) == "-"

    def test_missing_num_id_returns_dash(self) -> None:
        from app.rag.extraction import _docx_list_prefix

        num_pr = MagicMock()
        num_pr.find.return_value = None  # num_id_el is None
        assert _docx_list_prefix(MagicMock(), num_pr, {}) == "-"


class TestDocxTableToMd:
    def _cell(self, text: str) -> MagicMock:
        c = MagicMock()
        c.text = text
        return c

    def _make_table(self, rows: list[list[MagicMock]]) -> MagicMock:
        mock_table = MagicMock()
        mock_rows = []
        for row_cells in rows:
            mock_row = MagicMock()
            mock_row.cells = row_cells
            mock_rows.append(mock_row)
        mock_table.rows = mock_rows
        return mock_table

    def test_header_and_separator(self) -> None:
        tbl = self._make_table([[self._cell("Col A"), self._cell("Col B")], [self._cell("r1c1"), self._cell("r1c2")]])
        md = _docx_table_to_md(tbl)
        assert "| Col A | Col B |" in md
        assert "| --- | --- |" in md

    def test_body_rows_included(self) -> None:
        tbl = self._make_table(
            [
                [self._cell("H1"), self._cell("H2")],
                [self._cell("v1"), self._cell("v2")],
                [self._cell("v3"), self._cell("v4")],
            ]
        )
        md = _docx_table_to_md(tbl)
        assert "| v1 | v2 |" in md
        assert "| v3 | v4 |" in md

    def test_empty_table_returns_empty_string(self) -> None:
        mock_table = MagicMock()
        mock_table.rows = []
        assert _docx_table_to_md(mock_table) == ""

    def test_horizontally_merged_cell_appears_once(self) -> None:
        merged = self._cell("Merged Header")
        tbl = self._make_table(
            [
                [merged, merged, self._cell("Col B")],
                [self._cell("r1c1"), self._cell("r1c2"), self._cell("r1c3")],
            ]
        )
        assert _docx_table_to_md(tbl).count("Merged Header") == 1

    def test_two_distinct_cells_with_same_text_both_appear(self) -> None:
        a, b = self._cell("Same"), self._cell("Same")
        tbl = self._make_table([[a, b], [self._cell("v1"), self._cell("v2")]])
        header_line = _docx_table_to_md(tbl).splitlines()[0]
        assert header_line.count("Same") == 2

    def test_vertically_merged_cell_not_duplicated(self) -> None:
        v = self._cell("Tall Cell")
        tbl = self._make_table(
            [
                [v, self._cell("R1C2")],
                [v, self._cell("R2C2")],
                [self._cell("R3C1"), self._cell("R3C2")],
            ]
        )
        assert _docx_table_to_md(tbl).count("Tall Cell") == 1

    def test_fully_merged_row_dropped_not_emitted_blank(self) -> None:
        shared_a, shared_b = self._cell("A"), self._cell("B")
        tbl = self._make_table(
            [
                [shared_a, shared_b],
                [shared_a, shared_b],
            ]
        )
        lines = [line for line in _docx_table_to_md(tbl).splitlines() if line.strip()]
        assert len(lines) == 2, f"Unexpected lines: {lines}"

    def test_both_axes_merged(self) -> None:
        corner = self._cell("Corner")
        tbl = self._make_table(
            [
                [corner, corner, self._cell("Top Right")],
                [corner, corner, self._cell("Bot Right")],
            ]
        )
        md = _docx_table_to_md(tbl)
        assert md.count("Corner") == 1
        assert "Top Right" in md
        assert "Bot Right" in md


class TestBsTableToMd:
    def _make_table(self, html: str) -> Tag:
        from bs4 import BeautifulSoup, Tag

        soup = BeautifulSoup(html, "html.parser")
        tag = soup.find("table")
        assert isinstance(tag, Tag)
        return tag

    def test_header_row(self) -> None:
        tbl = self._make_table("<table><tr><th>A</th><th>B</th></tr></table>")
        assert "| A | B |" in _bs_table_to_md(tbl)

    def test_separator_row(self) -> None:
        tbl = self._make_table("<table><tr><th>A</th><th>B</th></tr></table>")
        assert "| --- | --- |" in _bs_table_to_md(tbl)

    def test_body_row(self) -> None:
        tbl = self._make_table("<table><tr><th>A</th></tr><tr><td>val</td></tr></table>")
        assert "| val |" in _bs_table_to_md(tbl)

    def test_empty_table_returns_empty_string(self) -> None:
        from bs4 import BeautifulSoup, Tag

        soup = BeautifulSoup("<table></table>", "html.parser")
        tag = soup.find("table")
        assert isinstance(tag, Tag)
        assert _bs_table_to_md(tag) == ""


class TestExtractToMarkdown:
    @pytest.mark.parametrize(
        "filename,expected_type",
        [
            ("notes.txt", DocType.TXT),
            ("notes.md", DocType.TXT),
            ("page.html", DocType.HTML),
            ("page.htm", DocType.HTML),
        ],
    )
    def test_routing_by_extension(self, filename: Any, expected_type: Any) -> None:
        data = MINIMAL_HTML if "html" in filename or "htm" in filename else TXT_CONTENT
        assert extract_to_markdown(data, filename).doc_type == expected_type

    def test_unsupported_extension_raises_value_error(self) -> None:
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_to_markdown(b"data", "file.pptx")

    def test_source_name_preserved(self) -> None:
        assert extract_to_markdown(TXT_CONTENT, "my_lecture.txt").source_name == "my_lecture.txt"

    def test_pdf_routing_calls_pymupdf(self) -> None:
        with (
            patch("app.rag.extraction.fitz.open"),
            patch("app.rag.extraction.pymupdf4llm.to_markdown", return_value="# MD") as m,
        ):
            extract_to_markdown(b"%PDF", "deck.pdf")
        m.assert_called_once()

    def test_docx_routing(self) -> None:
        from docx import Document

        doc = Document()
        doc.add_paragraph("hello")
        buf = io.BytesIO()
        doc.save(buf)
        assert extract_to_markdown(buf.getvalue(), "file.docx").doc_type == DocType.DOCX


class TestAllowedExtensionsFilter:
    def test_extension_not_in_allowed_list_raises(self) -> None:
        with patch("app.rag.extraction.get_settings") as mock:
            mock.return_value.RAG_ALLOWED_EXTENSIONS = "pdf"
            with pytest.raises(ValueError, match="not enabled for RAG processing"):
                extract_to_markdown(MINIMAL_HTML, "page.html")

    def test_allowed_extension_proceeds(self) -> None:
        with patch("app.rag.extraction.get_settings") as mock:
            mock.return_value.RAG_ALLOWED_EXTENSIONS = "txt"
            result = extract_to_markdown(TXT_CONTENT, "lecture.txt")
            assert result.doc_type == DocType.TXT

    def test_error_message_includes_filename_and_accepted_types(self) -> None:
        with patch("app.rag.extraction.get_settings") as mock:
            mock.return_value.RAG_ALLOWED_EXTENSIONS = "pdf,docx"
            with pytest.raises(ValueError, match=r"notes\.txt") as exc_info:
                extract_to_markdown(TXT_CONTENT, "notes.txt")
            msg = str(exc_info.value)
            assert ".pdf" in msg
            assert ".docx" in msg

    def test_empty_string_permits_all_supported(self) -> None:
        with patch("app.rag.extraction.get_settings") as mock:
            mock.return_value.RAG_ALLOWED_EXTENSIONS = ""
            result = extract_to_markdown(TXT_CONTENT, "lecture.txt")
            assert result.doc_type == DocType.TXT

    def test_uppercase_file_extension_normalised(self) -> None:
        with patch("app.rag.extraction.get_settings") as mock:
            mock.return_value.RAG_ALLOWED_EXTENSIONS = "txt"
            result = extract_to_markdown(TXT_CONTENT, "lecture.TXT")
            assert result.doc_type == DocType.TXT

    def test_disallowed_extension_raises_descriptive_error_not_unsupported(self) -> None:
        with patch("app.rag.extraction.get_settings") as mock:
            mock.return_value.RAG_ALLOWED_EXTENSIONS = "pdf"
            with pytest.raises(ValueError, match="not enabled for RAG processing"):
                extract_to_markdown(TXT_CONTENT, "notes.txt")


class TestSupportedExtensions:
    def test_exported_constant_exists(self) -> None:
        from app.rag.extraction import SUPPORTED_EXTENSIONS

        assert isinstance(SUPPORTED_EXTENSIONS, frozenset)

    def test_matches_dispatch_keys(self) -> None:
        from app.rag.extraction import _DISPATCH, SUPPORTED_EXTENSIONS

        assert SUPPORTED_EXTENSIONS == frozenset(_DISPATCH)

    def test_utils_uses_extraction_constant(self) -> None:
        from app.core_plugins.googledrive.utils import _is_supported_for_rag
        from app.rag.extraction import SUPPORTED_EXTENSIONS

        for ext in SUPPORTED_EXTENSIONS:
            assert _is_supported_for_rag(f"file.{ext}") is True
