"""
Integration tests for Layer 1 (extraction) → Layer 2 (chunking) pipeline.
Tests use realistic in-memory documents across all supported formats.
"""

import io

import pytest

from app.rag.chunking import chunk_document
from app.rag.extraction import extract_to_markdown
from app.rag.types import Chunk


def _run(data: bytes, filename: str) -> list[Chunk]:
    """Run the full Layer 1 → Layer 2 pipeline and return chunks."""
    result = extract_to_markdown(data, filename)
    return chunk_document(result)


def _all_content(chunks: list[Chunk]) -> str:
    return "\n".join(c.content for c in chunks)


STRUCTURED_HTML = b"""
<html><body>
  <div>
    <h1>Machine Learning Fundamentals</h1>
    <p>This course covers the core concepts of machine learning.</p>
    <div>
      <h2>Supervised Learning</h2>
      <p>Supervised learning uses labelled training data.</p>
      <h3>Linear Regression</h3>
      <p>Linear regression models a continuous output variable.</p>
      <h3>Classification</h3>
      <p>Classification predicts discrete class labels.</p>
    </div>
    <div>
      <h2>Unsupervised Learning</h2>
      <p>Unsupervised learning finds structure in unlabelled data.</p>
      <h3>Clustering</h3>
      <p>Clustering groups similar data points together.</p>
    </div>
  </div>
</body></html>
"""


class TestHTMLExtractionToChunking:
    def test_pipeline_returns_chunks(self) -> None:
        chunks = _run(STRUCTURED_HTML, "course.html")
        assert len(chunks) > 0

    def test_all_chunks_have_content(self) -> None:
        chunks = _run(STRUCTURED_HTML, "course.html")
        assert all(c.content.strip() for c in chunks)

    def test_chapter_metadata_extracted(self) -> None:
        chunks = _run(STRUCTURED_HTML, "course.html")
        chapters = [c.metadata.chapter for c in chunks if c.metadata.chapter]
        assert any("Machine Learning Fundamentals" in ch for ch in chapters)

    def test_section_metadata_extracted(self) -> None:
        chunks = _run(STRUCTURED_HTML, "course.html")
        sections = [c.metadata.section for c in chunks if c.metadata.section]
        assert any("Supervised Learning" in s for s in sections)
        assert any("Unsupervised Learning" in s for s in sections)

    def test_subsection_metadata_extracted(self) -> None:
        chunks = _run(STRUCTURED_HTML, "course.html")
        subsections = [c.metadata.subsection for c in chunks if c.metadata.subsection]
        assert any("Linear Regression" in s for s in subsections)
        assert any("Clustering" in s for s in subsections)

    def test_source_name_on_all_chunks(self) -> None:
        chunks = _run(STRUCTURED_HTML, "course.html")
        assert all(c.metadata.source_name == "course.html" for c in chunks)

    def test_body_text_present_in_chunks(self) -> None:
        chunks = _run(STRUCTURED_HTML, "course.html")
        content = _all_content(chunks)
        assert "labelled training data" in content
        assert "continuous output variable" in content

    def test_nested_divs_not_dropped(self) -> None:
        chunks = _run(STRUCTURED_HTML, "course.html")
        content = _all_content(chunks)
        assert "Classification" in content
        assert "Clustering" in content


STRUCTURED_MD = b"""\
# Introduction to Deep Learning

Deep learning is a subfield of machine learning.

## Neural Networks

Neural networks are inspired by the human brain.

### Backpropagation

Backpropagation is the algorithm used to train neural networks.

## Convolutional Networks

CNNs are primarily used for image recognition tasks.
"""


FLAT_MD = b"No headings here.\nJust plain paragraphs of text.\n"


class TestMarkdownExtractionToChunking:
    def test_structured_md_produces_multiple_chunks(self) -> None:
        chunks = _run(STRUCTURED_MD, "notes.md")
        assert len(chunks) > 1

    def test_chapter_metadata_present(self) -> None:
        chunks = _run(STRUCTURED_MD, "notes.md")
        chapters = [c.metadata.chapter for c in chunks if c.metadata.chapter]
        assert any("Introduction to Deep Learning" in ch for ch in chapters)

    def test_section_metadata_present(self) -> None:
        chunks = _run(STRUCTURED_MD, "notes.md")
        sections = [c.metadata.section for c in chunks if c.metadata.section]
        assert any("Neural Networks" in s for s in sections)
        assert any("Convolutional Networks" in s for s in sections)

    def test_subsection_metadata_present(self) -> None:
        chunks = _run(STRUCTURED_MD, "notes.md")
        subsections = [c.metadata.subsection for c in chunks if c.metadata.subsection]
        assert any("Backpropagation" in s for s in subsections)

    def test_flat_md_returns_single_chunk(self) -> None:
        chunks = _run(FLAT_MD, "flat.md")
        assert len(chunks) == 1
        assert chunks[0].metadata.chapter is None

    def test_txt_extension_also_works(self) -> None:
        chunks = _run(STRUCTURED_MD, "notes.txt")
        assert len(chunks) > 1

    def test_source_name_preserved(self) -> None:
        chunks = _run(STRUCTURED_MD, "notes.md")
        assert all(c.metadata.source_name == "notes.md" for c in chunks)


def _make_docx(sections: list[tuple[int, str]]) -> bytes:
    """
    Build a minimal .docx in memory.
    sections: list of (heading_level, text) — level 0 means paragraph.
    """
    from docx import Document

    doc = Document()
    for level, text in sections:
        if level == 0:
            doc.add_paragraph(text)
        else:
            doc.add_heading(text, level=level)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDOCXExtractionToChunking:
    @pytest.fixture()
    def structured_docx(self) -> bytes:
        return _make_docx(
            [
                (1, "Introduction"),
                (0, "This section introduces the topic."),
                (2, "Background"),
                (0, "Historical context is provided here."),
                (3, "Early Research"),
                (0, "Early research dates back to the 1950s."),
                (2, "Scope"),
                (0, "The scope of this course is broad."),
            ]
        )

    def test_pipeline_returns_chunks(self, structured_docx: bytes) -> None:
        chunks = _run(structured_docx, "lecture.docx")
        assert len(chunks) > 0

    def test_chapter_extracted(self, structured_docx: bytes) -> None:
        chunks = _run(structured_docx, "lecture.docx")
        chapters = [c.metadata.chapter for c in chunks if c.metadata.chapter]
        assert any("Introduction" in ch for ch in chapters)

    def test_section_extracted(self, structured_docx: bytes) -> None:
        chunks = _run(structured_docx, "lecture.docx")
        sections = [c.metadata.section for c in chunks if c.metadata.section]
        assert any("Background" in s for s in sections)
        assert any("Scope" in s for s in sections)

    def test_subsection_extracted(self, structured_docx: bytes) -> None:
        chunks = _run(structured_docx, "lecture.docx")
        subsections = [c.metadata.subsection for c in chunks if c.metadata.subsection]
        assert any("Early Research" in s for s in subsections)

    def test_body_text_present(self, structured_docx: bytes) -> None:
        chunks = _run(structured_docx, "lecture.docx")
        content = _all_content(chunks)
        assert "Historical context" in content
        assert "1950s" in content

    def test_source_name_preserved(self, structured_docx: bytes) -> None:
        chunks = _run(structured_docx, "lecture.docx")
        assert all(c.metadata.source_name == "lecture.docx" for c in chunks)


class TestPDFExtractionToChunking:
    @pytest.fixture()
    def structured_pdf_bytes(self) -> bytes:
        """
        Generate a minimal real PDF with heading-like bold text using pymupdf.
        pymupdf4llm infers headings from font size/weight, so we use a large
        font for headings and normal size for body text.
        """
        import pymupdf

        doc = pymupdf.open()  # type: ignore[no-untyped-call]
        page = doc.new_page()

        # H1-like: large font
        page.insert_text((50, 80), "Introduction to Reinforcement Learning", fontsize=20, fontname="helv")
        # Body
        page.insert_text((50, 120), "Reinforcement learning trains agents via rewards.", fontsize=11, fontname="helv")
        # H2-like: medium font
        page.insert_text((50, 160), "Markov Decision Processes", fontsize=16, fontname="helv")
        # Body
        page.insert_text((50, 200), "MDPs model sequential decision-making problems.", fontsize=11, fontname="helv")
        doc_bytes: bytes = doc.tobytes()  # type: ignore[no-untyped-call]
        return doc_bytes

    def test_pipeline_does_not_raise(self, structured_pdf_bytes: bytes) -> None:
        chunks = _run(structured_pdf_bytes, "rl.pdf")
        assert isinstance(chunks, list)

    def test_returns_at_least_one_chunk(self, structured_pdf_bytes: bytes) -> None:
        chunks = _run(structured_pdf_bytes, "rl.pdf")
        assert len(chunks) >= 1

    def test_body_text_present_in_chunks(self, structured_pdf_bytes: bytes) -> None:
        chunks = _run(structured_pdf_bytes, "rl.pdf")
        content = _all_content(chunks)
        assert "rewards" in content or "Reinforcement" in content

    def test_source_name_preserved(self, structured_pdf_bytes: bytes) -> None:
        chunks = _run(structured_pdf_bytes, "rl.pdf")
        assert all(c.metadata.source_name == "rl.pdf" for c in chunks)


class TestChunkingDeterminism:
    """Same input must always produce identical chunks (required for hash-based dedup)."""

    def test_html_is_deterministic(self) -> None:
        a = _run(STRUCTURED_HTML, "course.html")
        b = _run(STRUCTURED_HTML, "course.html")
        assert a == b

    def test_markdown_is_deterministic(self) -> None:
        a = _run(STRUCTURED_MD, "notes.md")
        b = _run(STRUCTURED_MD, "notes.md")
        assert a == b

    def test_txt_is_deterministic(self) -> None:
        data = b"Some plain text content.\nAnother line.\n"
        a = _run(data, "doc.txt")
        b = _run(data, "doc.txt")
        assert a == b

    def test_docx_is_deterministic(self) -> None:
        docx_bytes = _make_docx(
            [
                (1, "Chapter"),
                (0, "Body text here."),
                (2, "Section"),
                (0, "Section body."),
            ]
        )
        a = _run(docx_bytes, "doc.docx")
        b = _run(docx_bytes, "doc.docx")
        assert a == b


class TestChunkingEdgeCases:
    _LONG_SENTENCE = "This is a sentence that contributes tokens to the chunk. " * 20
    _OVERSIZED_MD = "\n".join(
        [
            "# Chapter One",
            "",
            "## Long Section",
            "",
            "  ".join([_LONG_SENTENCE] * 6),
            "",
            "## Short Section",
            "",
            "A brief paragraph.",
            "",
        ]
    ).encode()

    def test_oversized_chunk_is_split(self) -> None:
        chunks = _run(self._OVERSIZED_MD, "big.md")
        long_section_chunks = [c for c in chunks if c.metadata.section == "Long Section"]
        assert len(long_section_chunks) > 1

    def test_secondary_split_inherits_chapter_metadata(self) -> None:
        chunks = _run(self._OVERSIZED_MD, "big.md")
        long_section_chunks = [c for c in chunks if c.metadata.section == "Long Section"]
        assert all(c.metadata.chapter == "Chapter One" for c in long_section_chunks)

    def test_secondary_split_inherits_section_metadata(self) -> None:
        chunks = _run(self._OVERSIZED_MD, "big.md")
        long_section_chunks = [c for c in chunks if c.metadata.section == "Long Section"]
        assert all(c.metadata.section == "Long Section" for c in long_section_chunks)

    def test_short_section_unaffected_by_secondary_split(self) -> None:
        chunks = _run(self._OVERSIZED_MD, "big.md")
        short_chunks = [c for c in chunks if c.metadata.section == "Short Section"]
        assert len(short_chunks) == 1
        assert "brief paragraph" in short_chunks[0].content

    def test_h4_not_a_split_boundary(self) -> None:
        md = b"# Chapter\n\n## Section\n\n#### Deep Heading\n\nBody under deep heading.\n"
        chunks = _run(md, "deep.md")
        assert all(c.metadata.subsection is None for c in chunks)
        content = _all_content(chunks)
        assert "Deep Heading" in content
        assert "Body under deep heading" in content

    def test_non_sequential_headings_h1_then_h3(self) -> None:
        md = b"# Top\n\nIntro text.\n\n### Sub without parent section\n\nSub body.\n"
        chunks = _run(md, "nonseq.md")
        sub_chunks = [c for c in chunks if c.metadata.subsection == "Sub without parent section"]
        assert len(sub_chunks) == 1
        assert sub_chunks[0].metadata.chapter == "Top"
        assert sub_chunks[0].metadata.section is None

    def test_multiple_h1s_produce_separate_chapters(self) -> None:
        md = b"# Alpha\n\nAlpha body.\n\n# Beta\n\nBeta body.\n"
        chunks = _run(md, "multi.md")
        chapters = {c.metadata.chapter for c in chunks if c.metadata.chapter}
        assert "Alpha" in chapters
        assert "Beta" in chapters

    def test_heading_only_section_not_empty(self) -> None:
        md = b"# Chapter\n\n## Section\n\n### Subsection\n\nActual content here.\n"
        chunks = _run(md, "headings.md")
        assert all(c.content.strip() for c in chunks)

    def test_whitespace_only_returns_empty(self) -> None:
        chunks = _run(b"   \n\n\t\n  ", "empty.md")
        assert chunks == []

    def test_unicode_content_preserved(self) -> None:
        md = "# \u0627\u0644\u0639\u0646\u0648\u0627\u0646\n\n\u0645\u062d\u062a\u0648\u0649 \u0627\u0644\u0646\u0635 \u0628\u0627\u0644\u0639\u0631\u0628\u064a\u0629.\n\n## \u6807\u9898\n\n\u4e2d\u6587\u5185\u5bb9\u3002\n".encode()
        chunks = _run(md, "unicode.md")
        content = _all_content(chunks)
        assert "\u0645\u062d\u062a\u0648\u0649" in content
        assert "\u4e2d\u6587\u5185\u5bb9" in content

    def test_unicode_metadata_preserved(self) -> None:
        md = "# \u0627\u0644\u0639\u0646\u0648\u0627\u0646\n\nBody.\n".encode()
        chunks = _run(md, "unicode.md")
        assert any(c.metadata.chapter == "\u0627\u0644\u0639\u0646\u0648\u0627\u0646" for c in chunks)

    def test_all_chunks_have_source_name(self) -> None:
        chunks = _run(self._OVERSIZED_MD, "big.md")
        assert all(c.metadata.source_name == "big.md" for c in chunks)

    def test_html_h4_not_a_split_boundary(self) -> None:
        html = b"""
        <html><body>
          <h1>Chapter</h1>
          <h2>Section</h2>
          <h4>Deep Heading</h4>
          <p>Paragraph under h4.</p>
        </body></html>
        """
        chunks = _run(html, "deep.html")
        assert all(c.metadata.subsection is None for c in chunks)
        content = _all_content(chunks)
        assert "Paragraph under h4" in content
